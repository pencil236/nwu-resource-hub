from datetime import UTC, datetime, timedelta
from io import BytesIO

from docx import Document

from app.core.security import hash_value
from app.db import SessionLocal
from app.models import EmailCode


def _register(client) -> dict[str, str]:
    with SessionLocal() as db:
        db.add(
            EmailCode(
                email="owner@school.edu.cn",
                code_hash=hash_value("654321"),
                expires_at=datetime.now(UTC) + timedelta(minutes=10),
            )
        )
        db.commit()
    response = client.post(
        "/api/auth/register",
        json={
            "email": "owner@school.edu.cn",
            "code": "654321",
            "password": "strong-password",
            "display_name": "Owner",
        },
    )
    return {"Authorization": f"Bearer {response.json()['access_token']}"}


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("高等数学期末复习", 0)
    document.add_paragraph("重点包括极限、导数、积分和常微分方程。适合期末复习。")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_upload_confirm_search_and_download(client):
    headers = _register(client)
    response = client.post(
        "/api/resources",
        headers=headers,
        data={
            "title": "高数期末重点笔记",
            "description": "期末复习提纲",
            "experience": "考前一周配合课后题使用",
            "course": "高等数学",
            "category": "笔记",
            "tags": "高数,期末",
            "rights_confirmed": "true",
        },
        files={
            "file": (
                "math.docx",
                _docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201, response.text
    resource = response.json()
    assert resource["status"] == "waiting_confirmation"
    assert resource["ai_summary"]

    response = client.post(f"/api/resources/{resource['id']}/confirm", headers=headers)
    assert response.status_code == 200
    assert response.json()["status"] == "published"

    response = client.get("/api/search", params={"q": "积分复习"}, headers=headers)
    assert response.status_code == 200
    assert response.json()[0]["resource"]["id"] == resource["id"]

    response = client.get(f"/api/resources/{resource['id']}/download", headers=headers)
    assert response.status_code == 200
    assert response.content

    ticket = client.post(f"/api/resources/{resource['id']}/download-ticket", headers=headers)
    assert ticket.status_code == 200
    signed_download = client.get(ticket.json()["url"])
    assert signed_download.status_code == 200
    assert signed_download.content


def test_upload_requires_authentication(client):
    response = client.get("/api/resources")
    assert response.status_code == 401
